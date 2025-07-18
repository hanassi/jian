from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

num = 5
translated_topic_1 = """


"""

# 새 문서 생성
doc = Document()
doc.add_heading(f"CMMC 모듈 {num} 번역본", 0)

# 본문 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = Pt(11)

# 내용 추가
doc.add_heading("주제 1: ", level=1)
doc.add_paragraph(translated_topic_1.strip())

# 파일 저장
file_path = f"./CMMC_Module_{num}_Korean_Translation.docx"
doc.save(file_path)
file_path
